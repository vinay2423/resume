from bs4 import BeautifulSoup
from docx import Document

# Load your HTML content
with open('Resume-online.html', 'r', encoding='utf-8') as file:
    html_content = file.read()

# Parse HTML
soup = BeautifulSoup(html_content, 'html.parser')

# Create a Word document
doc = Document()

# Add parsed content to Word
for element in soup.body.children:  # Iterate over body elements
    if element.name == 'p':
        doc.add_paragraph(element.get_text())
    elif element.name == 'h1':
        doc.add_heading(element.get_text(), level=1)
    elif element.name == 'h2':
        doc.add_heading(element.get_text(), level=2)
    elif element.name == 'ul':  # Unordered list
        for li in element.find_all('li'):
            doc.add_paragraph(li.get_text(), style='ListBullet')
    elif element.name == 'ol':  # Ordered list
        for li in element.find_all('li'):
            doc.add_paragraph(li.get_text(), style='ListNumber')

# Save the document
doc.save('output.docx')

